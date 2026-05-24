"""Tests for Word OOXML merge-aware grid reading."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from src.parsers.word_merge_reader import collect_word_merged_cell_evidence, table_to_merge_aware_grid
from src.parsers.word_parser import extract_word_document

SHUTOFF = (
    Path(__file__).resolve().parents[2]
    / "pm_sample_inputs"
    / "input"
    / "edited_Shutoff_Condition_Spec.docx"
)


def test_shutoff_merge_evidence_when_flag_on():
    if not SHUTOFF.exists():
        return
    data = extract_word_document(SHUTOFF, cfg={"features": {"word_merge_geometry": True}})
    evidence = data.get("merged_cell_evidence") or []
    assert len(evidence) >= 1
    assert evidence[0].get("kind") == "table_merged_region"


def test_shutoff_merge_flag_off_has_no_evidence():
    if not SHUTOFF.exists():
        return
    data = extract_word_document(SHUTOFF, cfg={"features": {"word_merge_geometry": False}})
    assert data.get("merged_cell_evidence") == []


def test_shutoff_shutdown_decision_expression_regression():
    """Merge evidence must not replace legacy grid used for nested condition columns."""
    if not SHUTOFF.exists():
        return
    data = extract_word_document(SHUTOFF, cfg={"features": {"word_merge_geometry": True}})
    block = next((b for b in data.get("logic_blocks") or [] if b.get("name") == "SHUTOFF_DECISION"), None)
    assert block is not None
    expr = str(block.get("raw_expression") or "")
    assert "HUY = OK" in expr
    assert "OK_SHUTOFF = 1" in expr
    assert expr.count("OR OR") < 2
    assert "(OR OR OR" not in expr


def test_merge_flag_on_preserves_logic_block_count():
    if not SHUTOFF.exists():
        return
    off = extract_word_document(SHUTOFF, cfg={"features": {"word_merge_geometry": False}})
    on = extract_word_document(SHUTOFF, cfg={"features": {"word_merge_geometry": True}})
    assert len(off.get("logic_blocks") or []) == len(on.get("logic_blocks") or [])


def test_table_to_merge_aware_grid_returns_rows():
    if not SHUTOFF.exists():
        return
    doc = Document(str(SHUTOFF))
    grid, _ = table_to_merge_aware_grid(doc.tables[0])
    assert len(grid) >= 2
    refs = collect_word_merged_cell_evidence(doc.tables[0], SHUTOFF.name, table_id="table_1")
    assert isinstance(refs, list)
