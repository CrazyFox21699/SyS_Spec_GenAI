"""Word document extraction: signals, logic tables, formulas, transitions."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from src.engine.two_column_logic_parser import (
    build_alias_map,
    extract_footnote_refs,
    parse_table_to_logic_block,
)
from src.parsers.diagram_parser import extract_diagram_transitions
from src.parsers.ocr_local import analyze_docx_embedded_images
from src.parsers.paragraph_extractor import extract_from_paragraphs, link_footnotes
from src.parsers.table_logic_parser import (
    LogicBlock,
    extract_formulas_from_text,
    parse_transition_table,
    rows_from_grid,
    transitions_from_logic_blocks,
)
from src.classifiers.logic_spec_classifier import classify_logic_spec
from src.engine.memory_semantics_parser import parse_retention_rules
from src.engine.state_grammar_parser import (
    parse_state_blocks_from_paragraphs,
    parse_state_blocks_from_tables,
)
from src.models.spec_profile import LogicZone, build_spec_profile
from src.utils.feature_flags import feature_enabled
from src.parsers.two_column_table_parser import parse_control_condition_grid, tables_to_dicts
from src.parsers.signal_table_parser import parse_signal_grid
from src.parsers.word_merge_reader import build_row_branch_groups, collect_word_merged_cell_evidence, table_to_merge_aware_grid
from src.parsers.word_section_router import (
    annotate_source_with_zone,
    build_word_section_map,
    zone_for_table,
)

_LOGIC_CONTROL_TOKENS = (
    "control",
    "event",
    "function",
    "item",
    "judgment",
    "signal",
    "permission",
    "prohibition",
)


def _is_two_column_logic_table(header: list[str]) -> bool:
    hdr_joined = " ".join(header)
    if "logic" in header and "condition" in hdr_joined:
        return False
    has_condition = "condition" in hdr_joined or " cond" in hdr_joined or header[-1:] == ["cond"]
    has_control = any(any(t in h for t in _LOGIC_CONTROL_TOKENS) for h in header if h)
    return bool(has_condition and has_control)


def _flat_fallback_logic_blocks(
    table_grids: list[list[list[str]]],
    *,
    file_name: str,
    src_base: dict[str, Any],
    skip_indices: set[int],
) -> list[dict[str, Any]]:
    """When header routing finds nothing, emit one leaf block per condition-like row."""
    blocks: list[dict[str, Any]] = []
    seq = 0
    for ti, grid in enumerate(table_grids):
        if ti in skip_indices or len(grid) < 2:
            continue
        for ri, cells in enumerate(grid[1:], start=2):
            if len(cells) < 2:
                continue
            label = (cells[0] or "").strip()
            condition = (cells[1] or "").strip()
            if not condition or condition.upper() in {"AND", "OR", "NOT"}:
                continue
            if not label:
                label = f"Row {ri}"
            seq += 1
            block_id = f"WD_FLAT_{ti + 1}_{seq:03d}"
            blocks.append(
                {
                    "id": block_id,
                    "name": label[:120],
                    "raw_expression": condition,
                    "tree": {
                        "type": "leaf",
                        "signal": label[:80],
                        "value": condition,
                        "parse_status": "partial",
                    },
                    "block_type": "flat_fallback",
                    "parse_status": "partial",
                    "review_required": True,
                    "source": {**src_base, "table": f"table_{ti + 1}", "row": ri, "kind": "flat_fallback"},
                }
            )
    return blocks


def peek_word_text(path: Path, max_chars: int = 8000) -> str:
    try:
        doc = Document(str(path))
    except PackageNotFoundError:
        return ""
    parts: list[str] = []
    for p in doc.paragraphs[:200]:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables[:30]:
        for row in table.rows[:40]:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)[:max_chars]


def _table_to_grid(table, cfg: dict[str, Any] | None = None) -> list[list[str]]:
    if cfg is None:
        try:
            from src.utils.config_path import get_config_path
            from src.utils.yaml_utils import load_yaml

            cfg = load_yaml(get_config_path())
        except OSError:
            cfg = {}
    if feature_enabled(cfg, "word_merge_geometry", default=True):
        try:
            grid, _ = table_to_merge_aware_grid(table)
            if grid:
                return grid
        except Exception:
            pass
    return [[c.text.strip() for c in row.cells] for row in table.rows]


def extract_word_document(path: Path, *, cfg: dict[str, Any] | None = None) -> dict[str, Any]:
    """
    Full Word extraction: signals, condition definitions, logic blocks (merged tables),
    paragraph formulas, state transitions, embedded test reference rows.
    """
    try:
        doc = Document(str(path))
    except PackageNotFoundError as exc:
        return {
            "error": "invalid_docx_package",
            "message": str(exc),
            "signals": [],
            "logic_blocks": [],
            "transitions": [],
            "condition_definitions": [],
            "test_reference_rows": [],
            "two_column_tables": [],
            "alias_map": [],
            "footnote_definitions": [],
        }

    cfg = cfg or {}
    use_section_router = feature_enabled(cfg, "word_section_router", default=True)
    use_merge_geometry = feature_enabled(cfg, "word_merge_geometry", default=True)
    use_state_grammar = feature_enabled(cfg, "state_grammar_parser", default=True)
    use_memory_semantics = feature_enabled(cfg, "memory_semantics_parser", default=True)

    src_base = {"file": path.name, "document": path.name}
    section_map = build_word_section_map(doc) if use_section_router else {}
    table_grids: list[list[list[str]]] = []
    merged_cell_evidence: list[dict[str, Any]] = []

    def _source_tbl(table_index: int) -> dict[str, Any]:
        base = {**src_base, "table": f"table_{table_index + 1}"}
        if not use_section_router:
            return base
        zone = zone_for_table(table_index, section_map)
        return annotate_source_with_zone(base, zone)

    signals: list[dict[str, Any]] = []
    logic_blocks: list[LogicBlock] = []
    two_column_logic_dicts: list[dict[str, Any]] = []
    parsed_tc_tables: list[Any] = []
    transitions: list[dict[str, Any]] = []
    condition_definitions: list[dict[str, Any]] = []
    test_reference_rows: list[dict[str, Any]] = []
    full_text_parts: list[str] = []
    state_table_payloads: list[dict[str, Any]] = []

    header_keywords = ("signal", "name", "interface", "direction", "sender", "receiver", "initial", "fail")
    matched_table_indices: set[int] = set()

    for ti, table in enumerate(doc.tables):
        grid = _table_to_grid(table, cfg=cfg)
        if use_merge_geometry:
            try:
                merged_cell_evidence.extend(
                    collect_word_merged_cell_evidence(
                        table,
                        path.name,
                        table_id=f"table_{ti + 1}",
                    )
                )
            except Exception:
                pass
        if not grid:
            continue
        header = [c.lower() for c in grid[0]]
        table_grids.append(grid)
        source_tbl = _source_tbl(ti)

        # Two-column Control/Event/Judgment/Signal + Condition (gates embedded in condition column(s))
        hdr_joined = " ".join(header)
        if _is_two_column_logic_table(header):
            matched_table_indices.add(ti)
            state_table_payloads.append({"grid": grid, "source": source_tbl})
            merge_branch_by_row: dict[int, dict[str, Any]] = {}
            if use_merge_geometry:
                try:
                    cond_indices = [
                        i
                        for i, h in enumerate(header)
                        if "condition" in h or h == "cond" or i > 0
                    ]
                    merge_branch_by_row = build_row_branch_groups(
                        table, ctrl_idx=0, cond_indices=cond_indices[1:] if len(cond_indices) > 1 else cond_indices
                    )
                except Exception:
                    merge_branch_by_row = {}
            tc_parsed = parse_control_condition_grid(
                grid,
                source_tbl,
                table_id=f"T{ti+1}",
                merge_branch_by_row=merge_branch_by_row or None,
            )
            parsed_tc_tables.extend(tc_parsed)
            for pt in tc_parsed:
                if pt.table_kind == "logic":
                    two_column_logic_dicts.append(parse_table_to_logic_block(pt))
                elif pt.table_kind == "constant":
                    for row in pt.rows:
                        condition_definitions.append(
                            {
                                "name": row.control,
                                "definition": row.condition_raw,
                                "source": row.source,
                                "constant_value": row.parsed_hint,
                            }
                        )
            continue

        # Logic / control tables (separate Logic column)
        if "logic" in header and "condition" in header:
            matched_table_indices.add(ti)
            blocks = rows_from_grid(grid, source_tbl, block_id_prefix=f"WD{ti+1}")
            logic_blocks.extend(blocks)
            continue

        # Condition definition table
        if header[0] == "condition" and len(header) > 1 and "definition" in header[1]:
            matched_table_indices.add(ti)
            for ri, cells in enumerate(grid[1:], start=2):
                if len(cells) < 2 or not cells[0]:
                    continue
                condition_definitions.append(
                    {
                        "name": cells[0],
                        "definition": cells[1],
                        "source": {**source_tbl, "row": ri},
                    }
                )
            continue

        # Test reference (Given / When / Expected)
        if "given" in " ".join(header) and ("expected" in " ".join(header) or "when" in " ".join(header)):
            matched_table_indices.add(ti)
            hdr_map = {h: i for i, h in enumerate(header)}
            for ri, cells in enumerate(grid[1:], start=2):
                if not any(cells):
                    continue
                test_reference_rows.append(
                    {
                        "id": cells[hdr_map.get("no", 0)] if "no" in hdr_map else cells[0],
                        "given": cells[hdr_map.get("given", 1)] if "given" in hdr_map else "",
                        "when": cells[hdr_map.get("when", 2)] if "when" in hdr_map else "",
                        "expected": cells[hdr_map.get("expected", 3)] if "expected" in hdr_map else "",
                        "source": {**source_tbl, "row": ri},
                    }
                )
            continue

        # State transition outcome table (header may be Item | Expected Value)
        body_text_joined = " ".join(" ".join(r) for r in grid[1:6]).lower()
        if (
            "previous" in " ".join(header)
            or "next" in " ".join(header)
            or "previous state" in body_text_joined
        ):
            matched_table_indices.add(ti)
            transitions.extend(parse_transition_table(grid, source_tbl))
            continue

        # Signal-like tables
        if any(any(k in h for k in header_keywords) for h in header):
            matched_table_indices.add(ti)
            signals.extend(_extract_signal_rows_from_grid(grid, source_tbl))

    paragraph_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text_parts = paragraph_lines

    para_data = extract_from_paragraphs(paragraph_lines, path.name)
    condition_definitions.extend(para_data.get("condition_definitions", []))
    transitions.extend(para_data.get("transitions", []))

    body_text = "\n".join(full_text_parts)
    logic_blocks.extend(
        extract_formulas_from_text(body_text, {**src_base, "kind": "paragraph"})
    )

    # Transition narrative: "NORMAL → SHUT_OFF"
    for m in re.finditer(
        r"transition[s]?\s+(?:from\s+)?(\w+)\s*(?:→|->)\s*(\w+)",
        body_text,
        re.I,
    ):
        transitions.append(
            {
                "id": f"SM_P_{len(transitions)+1:03d}",
                "from_state": m.group(1).upper(),
                "to_state": m.group(2).upper(),
                "event": "narrative_transition",
                "raw_condition": m.group(0),
                "source": {**src_base, "kind": "paragraph"},
                "confidence": "medium",
                "review_required": True,
            }
        )

    logic_dicts = [_block_to_dict(b) for b in logic_blocks]
    logic_dicts.extend(two_column_logic_dicts)
    if not logic_dicts:
        logic_dicts.extend(
            _flat_fallback_logic_blocks(
                table_grids,
                file_name=path.name,
                src_base=src_base,
                skip_indices=matched_table_indices,
            )
        )
    transitions.extend(transitions_from_logic_blocks(logic_dicts))
    alias_map = build_alias_map(parsed_tc_tables)
    footnote_refs = extract_footnote_refs(parsed_tc_tables)
    footnotes = link_footnotes(
        footnote_refs,
        para_data.get("footnote_map", {}),
        condition_definitions,
    )
    diagram_transitions = extract_diagram_transitions(paragraph_lines, path.name)
    embedded_image_analysis = analyze_docx_embedded_images(path)
    code_definitions: list[dict[str, Any]] = list(para_data.get("code_definitions", []))
    state_rules: list[dict[str, Any]] = list(para_data.get("state_rules", []))
    for analysis in embedded_image_analysis:
        condition_definitions.extend(analysis.get("condition_definitions", []))
        transitions.extend(analysis.get("transitions", []))
        code_definitions.extend(analysis.get("code_definitions", []))
        state_rules.extend(analysis.get("state_rules", []))
    transitions.extend(diagram_transitions)

    classifier = classify_logic_spec(body_text, table_samples=table_grids)
    spec_profile = build_spec_profile(
        file_name=path.name,
        is_logic_spec=bool(classifier.get("is_logic_spec")),
        classifier_score=float(classifier.get("score") or 0.0),
        classifier_signals=list(classifier.get("signals") or []),
        section_zones=list(section_map.get("sections") or []) if use_section_router else [],
    )

    state_machines: list[dict[str, Any]] = []
    if use_state_grammar:
        state_machines.extend(parse_state_blocks_from_paragraphs(paragraph_lines, src_base))
        state_machines.extend(parse_state_blocks_from_tables(state_table_payloads))

    retention_rules: list[dict[str, Any]] = []
    if use_memory_semantics:
        retention_rules = parse_retention_rules(paragraph_lines, src_base)

    return {
        "signals": signals,
        "logic_blocks": logic_dicts,
        "transitions": transitions,
        "condition_definitions": condition_definitions,
        "test_reference_rows": test_reference_rows,
        "two_column_tables": tables_to_dicts(parsed_tc_tables),
        "alias_map": alias_map,
        "footnote_definitions": footnotes,
        "code_definitions": code_definitions,
        "state_rules": state_rules,
        "diagram_transitions": diagram_transitions,
        "embedded_image_analysis": embedded_image_analysis,
        "full_text_sample": body_text[:4000],
        "spec_profile": spec_profile,
        "word_section_map": section_map if use_section_router else {},
        "merged_cell_evidence": merged_cell_evidence,
        "state_machines": state_machines,
        "retention_rules": retention_rules,
    }


def _block_to_dict(b: LogicBlock) -> dict[str, Any]:
    return {
        "id": b.id,
        "name": b.name,
        "raw_expression": b.raw_expression,
        "tree": b.tree,
        "block_type": b.block_type,
        "parse_status": b.parse_status,
        "review_required": b.review_required,
        "source": b.source,
        "rows": [
            {"control": r.control, "logic": r.logic, "condition": r.condition, "detail": r.detail}
            for r in b.rows
        ],
    }


def _extract_signal_rows_from_grid(grid: list[list[str]], source: dict[str, Any]) -> list[dict[str, Any]]:
    return parse_signal_grid(grid, source)


def extract_word_signals(path: Path) -> list[dict[str, Any]]:
    """Backward-compatible: signals only."""
    return extract_word_document(path).get("signals", [])
